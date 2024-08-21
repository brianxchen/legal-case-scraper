from docx import Document
import csv
import os

# Some cells contain nested tables. Parse the nested tables to append them all 
def extract_cell_text(cell):
    cell_text = []
    if cell.tables:
        for nested_table in cell.tables:
            for row in nested_table.rows:
                row_text = ""
                for cell in row.cells:
                    row_text += cell.text
                
                # for some reason delimiter characters are saved in the word file as '\xa0', remove these
                cell_text.append(row_text.replace('\xa0', '')) 
    else:
        # If no table, just add the cell's text
        cell_text.append(cell.text.replace('\xa0', ''))
    return cell_text

data = []

files = []
current_directory = os.getcwd()

# Loop through all files in the directory
for filename in os.listdir(current_directory):
    if filename.endswith('.docx') and not filename.startswith('~$'):
        # Some strange temporary file memory issue if you don't remove the temp files 
        # that word makes when you open a file..........
        files.append(filename)

for filename in files:
    doc = Document(filename)
    for table in doc.tables:
        for row in table.rows:
            row_data = ['', '', '', '', '']
            for cell in row.cells:
                cell_texts = extract_cell_text(cell)
                for text in cell_texts:
                    cleaned_text = text.split(': ', 1)[-1] if ': ' in text else text
                    if "represented by" in cleaned_text:
                        continue
                    elif "Phone:" in cleaned_text:
                        row_data[2] = cleaned_text.replace("Phone:", "").strip()
                    elif "Fax:" in cleaned_text:
                        row_data[3] = cleaned_text.replace("Fax:", "").strip()
                    elif "Email:" in text:
                        row_data[4] = cleaned_text.replace("Email:", "").strip()
                    else:
                        if not row_data[0]:
                            # We only need to log if it's a plaintiff or defendant or whatever, not the actual name
                            row_data[0] = text.split(': ', 1)[0] if ':' in text else text
                        else:
                            row_data[1] = text
            
            data.append(row_data)

headers = ["Plaintiff/Defendant", "Attorney", "Phone", "Fax", "Email"]

# save to CSV
with open('output.csv', 'w', newline='', encoding='utf-8') as file:
    writer = csv.writer(file)
    writer.writerow(headers)
    writer.writerows(data)

print("Done parsing :)")
